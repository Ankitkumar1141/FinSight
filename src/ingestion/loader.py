import re
import logging
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".md"}

DOC_TYPE_KEYWORDS = {
    "10-k": "annual_report",
    "annual report": "annual_report",
    "10k": "annual_report",
    "earnings": "earnings_transcript",
    "transcript": "earnings_transcript",
    "investor presentation": "investor_presentation",
    "presentation": "investor_presentation",
    "press release": "financial_news",
    "news": "financial_news",
}


def detect_doc_type(filename: str, first_page_text: str = "") -> str:
    filename_raw   = filename.lower()
    # Normalised form: underscores/hyphens → spaces so "annual_report" matches "annual report"
    filename_norm  = filename_raw.replace("_", " ").replace("-", " ")
    text_lower     = first_page_text.lower()
    for keyword, doc_type in DOC_TYPE_KEYWORDS.items():
        # Check raw (preserves "10-k"), normalised (handles "annual_report"), and page text
        if keyword in filename_raw or keyword in filename_norm or keyword in text_lower:
            return doc_type
    return "unknown"


def extract_year(text: str) -> str:
    years = re.findall(r"\b(20\d{2})\b", text)
    if years:
        return Counter(years).most_common(1)[0][0]
    return "unknown"


def extract_company_name(filename: str) -> str:
    # Try common patterns like "AAPL_10K_2023" or "Apple_Annual_Report"
    parts = Path(filename).stem.replace("-", "_").split("_")
    if parts:
        return parts[0].upper()
    return "unknown"


def load_pdf(file_path: str) -> List[Dict[str, Any]]:
    import pdfplumber

    pages = []
    with pdfplumber.open(file_path) as pdf:
        first_page_text = pdf.pages[0].extract_text() or "" if pdf.pages else ""
        filename = Path(file_path).name
        doc_type = detect_doc_type(filename, first_page_text)
        year = extract_year(first_page_text)
        company = extract_company_name(filename)

        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if not text.strip():
                continue

            # Extract tables and convert to readable text
            table_text = ""
            for table in page.extract_tables():
                if not table:
                    continue
                for row in table:
                    if row:
                        row_str = " | ".join(str(cell).strip() if cell else "" for cell in row)
                        table_text += row_str + "\n"

            full_text = text
            if table_text:
                full_text += "\n[TABLE DATA]\n" + table_text

            pages.append({
                "content": full_text,
                "metadata": {
                    "source": filename,
                    "page": i + 1,
                    "doc_type": doc_type,
                    "year": year,
                    "company": company,
                    "file_path": file_path,
                },
            })

    logger.info(f"Loaded {len(pages)} pages from PDF: {file_path}")
    return pages


def load_txt(file_path: str) -> List[Dict[str, Any]]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    filename = Path(file_path).name
    doc_type = detect_doc_type(filename, content[:500])
    year = extract_year(content[:1000])
    company = extract_company_name(filename)

    return [{
        "content": content,
        "metadata": {
            "source": filename,
            "page": 1,
            "doc_type": doc_type,
            "year": year,
            "company": company,
            "file_path": file_path,
        },
    }]


def load_docx(file_path: str) -> List[Dict[str, Any]]:
    import docx

    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(paragraphs)

    # Also extract tables — financial .docx files often store key data in tables
    table_lines = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                table_lines.append(row_text)
    if table_lines:
        content += "\n\n[TABLE DATA]\n" + "\n".join(table_lines)

    filename = Path(file_path).name
    doc_type = detect_doc_type(filename, content[:500])
    year = extract_year(content[:1000])
    company = extract_company_name(filename)

    return [{
        "content": content,
        "metadata": {
            "source": filename,
            "page": 1,
            "doc_type": doc_type,
            "year": year,
            "company": company,
            "file_path": file_path,
        },
    }]


def load_document(file_path: str) -> List[Dict[str, Any]]:
    ext = Path(file_path).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_EXTENSIONS}")

    loaders = {
        ".pdf": load_pdf,
        ".txt": load_txt,
        ".docx": load_docx,
        ".md": load_txt,
    }

    pages = loaders[ext](file_path)
    if not pages:
        raise ValueError(f"No content could be extracted from: {file_path}")
    return pages
